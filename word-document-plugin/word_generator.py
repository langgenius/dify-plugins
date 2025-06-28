"""
Dify Word Document Output Plugin
Generates Word documents from text content with formatting options
"""

import io
from typing import Any, Dict, List
from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage, ToolParameter, ToolParameterOption, I18nObject


class WordDocumentGeneratorTool(Tool):
    """
    Tool for generating Word documents from text content
    """
    
    def _invoke(self, user_id: str, tool_parameters: Dict[str, Any]) -> List[ToolInvokeMessage]:
        """
        Invoke the Word document generator tool
        
        Args:
            user_id: The user ID
            tool_parameters: Tool parameters including content, title, formatting options
            
        Returns:
            List[ToolInvokeMessage]: The generated Word document
        """
        try:
            # Extract parameters
            content = tool_parameters.get('content', '')
            title = tool_parameters.get('title', 'Document')
            font_name = tool_parameters.get('font_name', 'Calibri')
            font_size = tool_parameters.get('font_size', 11)
            include_header = tool_parameters.get('include_header', True)
            include_footer = tool_parameters.get('include_footer', True)
            alignment = tool_parameters.get('alignment', 'left')
            
            # Validate required parameters
            if not content:
                raise ValueError("Content is required")
            
            # Create Word document
            doc = Document()
            
            # Set document properties
            doc.core_properties.title = title
            doc.core_properties.author = "Dify AI Assistant"
            doc.core_properties.created = datetime.now()
            
            # Add header if requested
            if include_header:
                self._add_header(doc, title)
            
            # Add title
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing after title
            doc.add_paragraph()
            
            # Process content - split by paragraphs and handle formatting
            paragraphs = content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a heading (starts with #)
                    if para_text.strip().startswith('#'):
                        self._add_heading(doc, para_text.strip(), font_name)
                    # Check if it's a bullet point (starts with - or *)
                    elif para_text.strip().startswith(('-', '*')):
                        self._add_bullet_point(doc, para_text.strip(), font_name, font_size, alignment)
                    # Check if it's a numbered list (starts with digit.)
                    elif para_text.strip().split('.')[0].isdigit():
                        self._add_numbered_item(doc, para_text.strip(), font_name, font_size, alignment)
                    else:
                        # Regular paragraph
                        self._add_paragraph(doc, para_text.strip(), font_name, font_size, alignment)
            
            # Add footer if requested
            if include_footer:
                self._add_footer(doc)
            
            # Save document to bytes
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Get document bytes
            doc_bytes = doc_buffer.getvalue()
            
            # Generate filename
            filename = f"{title.replace(' ', '_')}.docx"
            
            return [
                ToolInvokeMessage(
                    type=ToolInvokeMessage.MessageType.TEXT,
                    message=ToolInvokeMessage.TextMessage(
                        text=f"Word document '{title}' has been generated successfully."
                    )
                ),
                ToolInvokeMessage(
                    type=ToolInvokeMessage.MessageType.BLOB,
                    message=ToolInvokeMessage.BlobMessage(
                        blob=doc_bytes
                    ),
                    meta={
                        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "filename": filename
                    }
                )
            ]
            
        except Exception as e:
            return [
                ToolInvokeMessage(
                    type=ToolInvokeMessage.MessageType.TEXT,
                    message=ToolInvokeMessage.TextMessage(
                        text=f"Error generating Word document: {str(e)}"
                    )
                )
            ]
    
    def _add_header(self, doc: DocumentType, title: str):
        """Add header to document"""
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = title
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_footer(self, doc: DocumentType):
        """Add footer to document"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated by Dify AI Assistant - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_heading(self, doc: DocumentType, text: str, font_name: str):
        """Add heading to document"""
        # Remove # symbols and determine heading level
        heading_level = 0
        clean_text = text
        while clean_text.startswith('#'):
            heading_level += 1
            clean_text = clean_text[1:].strip()
        
        # Limit heading level to 3
        heading_level = min(heading_level, 3)
        
        if heading_level > 0:
            heading = doc.add_heading(clean_text, level=heading_level)
            heading.runs[0].font.name = font_name
        else:
            # Fallback to bold paragraph if no # found
            para = doc.add_paragraph()
            run = para.add_run(clean_text)
            run.font.bold = True
            run.font.name = font_name
    
    def _add_bullet_point(self, doc: DocumentType, text: str, font_name: str, font_size: int, alignment: str):
        """Add bullet point to document"""
        # Remove bullet symbols
        clean_text = text.lstrip('-*').strip()
        
        para = doc.add_paragraph(clean_text, style='List Bullet')
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        self._set_alignment(para, alignment)
    
    def _add_numbered_item(self, doc: DocumentType, text: str, font_name: str, font_size: int, alignment: str):
        """Add numbered list item to document"""
        para = doc.add_paragraph(text, style='List Number')
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        self._set_alignment(para, alignment)
    
    def _add_paragraph(self, doc: DocumentType, text: str, font_name: str, font_size: int, alignment: str):
        """Add regular paragraph to document"""
        para = doc.add_paragraph(text)
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        self._set_alignment(para, alignment)
    
    def _set_alignment(self, paragraph, alignment: str):
        """Set paragraph alignment"""
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.alignment = alignment_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)
    
    def get_runtime_parameters(self) -> List[ToolParameter]:
        """
        Return a list of runtime parameters for the tool.
        """
        return [
            ToolParameter(
                name="content",
                label=I18nObject(en_US="Content"),
                type=ToolParameter.ToolParameterType.STRING,
                required=True,
                human_description=I18nObject(en_US="The text content to include in the Word document. Supports basic formatting with # for headings, - or * for bullets, and numbered lists."),
                form=ToolParameter.ToolParameterForm.FORM
            ),
            ToolParameter(
                name="title",
                label=I18nObject(en_US="Document Title"),
                type=ToolParameter.ToolParameterType.STRING,
                required=False,
                default="Document",
                human_description=I18nObject(en_US="The title of the document"),
                form=ToolParameter.ToolParameterForm.FORM
            ),
            ToolParameter(
                name="font_name", 
                label=I18nObject(en_US="Font Name"),
                type=ToolParameter.ToolParameterType.SELECT,
                required=False,
                default="Calibri",
                options=[
                    ToolParameterOption(label=I18nObject(en_US="Calibri"), value="Calibri"),
                    ToolParameterOption(label=I18nObject(en_US="Arial"), value="Arial"),
                    ToolParameterOption(label=I18nObject(en_US="Times New Roman"), value="Times New Roman"),
                    ToolParameterOption(label=I18nObject(en_US="Helvetica"), value="Helvetica"),
                    ToolParameterOption(label=I18nObject(en_US="Georgia"), value="Georgia")
                ],
                human_description=I18nObject(en_US="Font family for the document text"),
                form=ToolParameter.ToolParameterForm.FORM
            ),
            ToolParameter(
                name="font_size",
                label=I18nObject(en_US="Font Size"),
                type=ToolParameter.ToolParameterType.NUMBER,
                required=False,
                default=11,
                human_description=I18nObject(en_US="Font size for the document text (8-72)"),
                form=ToolParameter.ToolParameterForm.FORM
            ),
            ToolParameter(
                name="alignment",
                label=I18nObject(en_US="Text Alignment"),
                type=ToolParameter.ToolParameterType.SELECT,
                required=False,
                default="left",
                options=[
                    ToolParameterOption(label=I18nObject(en_US="Left"), value="left"),
                    ToolParameterOption(label=I18nObject(en_US="Center"), value="center"),
                    ToolParameterOption(label=I18nObject(en_US="Right"), value="right"),
                    ToolParameterOption(label=I18nObject(en_US="Justify"), value="justify")
                ],
                human_description=I18nObject(en_US="Text alignment for paragraphs"),
                form=ToolParameter.ToolParameterForm.FORM
            ),
            ToolParameter(
                name="include_header",
                label=I18nObject(en_US="Include Header"),
                type=ToolParameter.ToolParameterType.BOOLEAN,
                required=False,
                default=True,
                human_description=I18nObject(en_US="Include document title in header"),
                form=ToolParameter.ToolParameterForm.FORM
            ),
            ToolParameter(
                name="include_footer",
                label=I18nObject(en_US="Include Footer"),
                type=ToolParameter.ToolParameterType.BOOLEAN,
                required=False,
                default=True,
                human_description=I18nObject(en_US="Include timestamp footer"),
                form=ToolParameter.ToolParameterForm.FORM
            )
        ]